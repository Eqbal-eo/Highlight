#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
برنامج استخراج النصوص المحددة باللون الأصفر من ملفات PDF
PDF Highlight Text Extractor

يقوم هذا البرنامج باستخراج النصوص المحددة (highlighted) باللون الأصفر من ملفات PDF
وحفظها في ملف نصي أو Word
"""

import fitz  # PyMuPDF
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
import os
from datetime import datetime
from docx import Document
from docx.shared import Inches


class PDFHighlightExtractor:
    def __init__(self):
        self.root = tk.Tk()
        self.root.title("مستخرج النصوص المحددة من PDF")
        self.root.geometry("600x500")
        self.root.configure(bg='#f0f0f0')
        
        # متغيرات
        self.pdf_file = None
        self.extracted_highlights = []
        
        self.setup_ui()
    
    def setup_ui(self):
        """إعداد واجهة المستخدم"""
        # العنوان الرئيسي
        title_label = tk.Label(
            self.root, 
            text="مستخرج النصوص المحددة من PDF",
            font=("Arial", 16, "bold"),
            bg='#f0f0f0',
            fg='#2c3e50'
        )
        title_label.pack(pady=20)
        
        # إطار اختيار الملف
        file_frame = tk.Frame(self.root, bg='#f0f0f0')
        file_frame.pack(pady=10, padx=20, fill='x')
        
        tk.Label(
            file_frame, 
            text="اختر ملف PDF:",
            font=("Arial", 12),
            bg='#f0f0f0'
        ).pack(anchor='w')
        
        self.file_path_var = tk.StringVar()
        self.file_entry = tk.Entry(
            file_frame, 
            textvariable=self.file_path_var,
            font=("Arial", 10),
            width=50,
            state='readonly'
        )
        self.file_entry.pack(side='left', padx=(0, 10), fill='x', expand=True)
        
        browse_btn = tk.Button(
            file_frame,
            text="تصفح",
            command=self.browse_file,
            bg='#3498db',
            fg='white',
            font=("Arial", 10),
            width=10
        )
        browse_btn.pack(side='right')
        
        # زر الاستخراج
        extract_btn = tk.Button(
            self.root,
            text="استخراج النصوص المحددة",
            command=self.extract_highlights,
            bg='#e74c3c',
            fg='white',
            font=("Arial", 12, "bold"),
            pady=10
        )
        extract_btn.pack(pady=20)
        
        # منطقة عرض النتائج
        results_frame = tk.Frame(self.root, bg='#f0f0f0')
        results_frame.pack(pady=10, padx=20, fill='both', expand=True)
        
        tk.Label(
            results_frame,
            text="النصوص المستخرجة:",
            font=("Arial", 12, "bold"),
            bg='#f0f0f0'
        ).pack(anchor='w')
        
        # منطقة النص مع شريط التمرير
        text_frame = tk.Frame(results_frame)
        text_frame.pack(fill='both', expand=True)
        
        self.results_text = tk.Text(
            text_frame,
            font=("Arial", 10),
            wrap='word',
            bg='white',
            relief='sunken',
            borderwidth=1
        )
        
        scrollbar = tk.Scrollbar(text_frame)
        scrollbar.pack(side='right', fill='y')
        self.results_text.pack(side='left', fill='both', expand=True)
        
        self.results_text.config(yscrollcommand=scrollbar.set)
        scrollbar.config(command=self.results_text.yview)
        
        # إطار أزرار الحفظ
        save_frame = tk.Frame(self.root, bg='#f0f0f0')
        save_frame.pack(pady=10)
        
        save_txt_btn = tk.Button(
            save_frame,
            text="حفظ كملف نصي",
            command=self.save_as_txt,
            bg='#27ae60',
            fg='white',
            font=("Arial", 10),
            width=15
        )
        save_txt_btn.pack(side='left', padx=5)
        
        save_docx_btn = tk.Button(
            save_frame,
            text="حفظ كملف Word",
            command=self.save_as_docx,
            bg='#8e44ad',
            fg='white',
            font=("Arial", 10),
            width=15
        )
        save_docx_btn.pack(side='left', padx=5)
        
        # شريط الحالة
        self.status_var = tk.StringVar()
        self.status_var.set("جاهز...")
        status_label = tk.Label(
            self.root,
            textvariable=self.status_var,
            relief='sunken',
            anchor='w',
            bg='#ecf0f1',
            font=("Arial", 9)
        )
        status_label.pack(side='bottom', fill='x')
    
    def browse_file(self):
        """تصفح واختيار ملف PDF"""
        file_path = filedialog.askopenfilename(
            title="اختر ملف PDF",
            filetypes=[("PDF files", "*.pdf"), ("All files", "*.*")]
        )
        
        if file_path:
            self.pdf_file = file_path
            self.file_path_var.set(file_path)
            self.status_var.set(f"تم اختيار الملف: {os.path.basename(file_path)}")
    
    def extract_highlights(self):
        """استخراج النصوص المحددة من PDF"""
        if not self.pdf_file:
            messagebox.showerror("خطأ", "يرجى اختيار ملف PDF أولاً!")
            return
        
        try:
            self.status_var.set("جاري استخراج النصوص المحددة...")
            self.root.update()
            
            # فتح ملف PDF
            doc = fitz.open(self.pdf_file)
            self.extracted_highlights = []
            
            # البحث في كل صفحة
            for page_num in range(len(doc)):
                page = doc[page_num]
                
                # البحث عن التحديدات (annotations)
                annotations = page.annots()
                
                for annot in annotations:
                    # التحقق من نوع التحديد
                    if annot.type[1] == 'Highlight':
                        # الحصول على النص المحدد
                        highlighted_text = self.get_highlighted_text(page, annot)
                        
                        if highlighted_text:
                            highlight_info = {
                                'page': page_num + 1,
                                'text': highlighted_text.strip(),
                                'color': self.get_annot_color(annot)
                            }
                            
                            # التحقق من اللون الأصفر
                            if self.is_yellow_highlight(highlight_info['color']):
                                self.extracted_highlights.append(highlight_info)
            
            doc.close()
            
            # عرض النتائج
            self.display_results()
            
            self.status_var.set(f"تم استخراج {len(self.extracted_highlights)} نص محدد باللون الأصفر")
            
        except Exception as e:
            messagebox.showerror("خطأ", f"حدث خطأ أثناء استخراج النصوص: {str(e)}")
            self.status_var.set("حدث خطأ أثناء الاستخراج")
    
    def get_highlighted_text(self, page, annot):
        """استخراج النص المحدد من التعليق التوضيحي"""
        try:
            # الحصول على مستطيل التحديد
            rect = annot.rect
            
            # استخراج النص من المنطقة المحددة
            text_instances = page.get_text("dict")
            highlighted_text = ""
            
            for block in text_instances["blocks"]:
                if "lines" in block:
                    for line in block["lines"]:
                        for span in line["spans"]:
                            span_rect = fitz.Rect(span["bbox"])
                            # التحقق من تداخل النص مع منطقة التحديد
                            if span_rect.intersects(rect):
                                highlighted_text += span["text"]
            
            # إذا لم نجد نص بالطريقة السابقة، نجرب طريقة أخرى
            if not highlighted_text.strip():
                highlighted_text = page.get_textbox(rect)
            
            return highlighted_text
            
        except Exception as e:
            print(f"خطأ في استخراج النص: {e}")
            return ""
    
    def get_annot_color(self, annot):
        """الحصول على لون التحديد"""
        try:
            # محاولة الحصول على اللون
            color = annot.colors.get("stroke", None)
            if not color:
                color = annot.colors.get("fill", None)
            
            return color if color else [1.0, 1.0, 0.0]  # أصفر افتراضي
        except:
            return [1.0, 1.0, 0.0]  # أصفر افتراضي
    
    def is_yellow_highlight(self, color):
        """التحقق من كون اللون أصفر"""
        if not color or len(color) < 3:
            return True  # افتراضياً نعتبره أصفر
        
        # التحقق من القيم RGB للون الأصفر
        # اللون الأصفر عادة ما يكون (1.0, 1.0, 0.0) أو قريب منه
        r, g, b = color[0], color[1], color[2]
        
        # نتحقق من أن الأحمر والأخضر مرتفعان والأزرق منخفض
        return r > 0.7 and g > 0.7 and b < 0.3
    
    def display_results(self):
        """عرض النتائج في منطقة النص"""
        self.results_text.delete(1.0, tk.END)
        
        if not self.extracted_highlights:
            self.results_text.insert(tk.END, "لم يتم العثور على نصوص محددة باللون الأصفر في هذا الملف.")
            return
        
        for i, highlight in enumerate(self.extracted_highlights, 1):
            self.results_text.insert(tk.END, f"[{i}] صفحة {highlight['page']}:\n")
            self.results_text.insert(tk.END, f"{highlight['text']}\n")
            self.results_text.insert(tk.END, "-" * 50 + "\n\n")
    
    def save_as_txt(self):
        """حفظ النتائج كملف نصي"""
        if not self.extracted_highlights:
            messagebox.showwarning("تحذير", "لا توجد نصوص لحفظها!")
            return
        
        file_path = filedialog.asksaveasfilename(
            defaultextension=".txt",
            filetypes=[("Text files", "*.txt"), ("All files", "*.*")],
            title="حفظ كملف نصي"
        )
        
        if file_path:
            try:
                with open(file_path, 'w', encoding='utf-8') as f:
                    f.write("النصوص المحددة باللون الأصفر من PDF\n")
                    f.write("=" * 50 + "\n")
                    f.write(f"الملف المصدر: {os.path.basename(self.pdf_file)}\n")
                    f.write(f"تاريخ الاستخراج: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
                    f.write(f"عدد النصوص المستخرجة: {len(self.extracted_highlights)}\n")
                    f.write("=" * 50 + "\n\n")
                    
                    for i, highlight in enumerate(self.extracted_highlights, 1):
                        f.write(f"[{i}] صفحة {highlight['page']}:\n")
                        f.write(f"{highlight['text']}\n")
                        f.write("-" * 50 + "\n\n")
                
                messagebox.showinfo("نجح", f"تم حفظ الملف بنجاح في:\n{file_path}")
                self.status_var.set("تم حفظ الملف النصي بنجاح")
                
            except Exception as e:
                messagebox.showerror("خطأ", f"حدث خطأ أثناء حفظ الملف: {str(e)}")
    
    def save_as_docx(self):
        """حفظ النتائج كملف Word"""
        if not self.extracted_highlights:
            messagebox.showwarning("تحذير", "لا توجد نصوص لحفظها!")
            return
        
        file_path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word documents", "*.docx"), ("All files", "*.*")],
            title="حفظ كملف Word"
        )
        
        if file_path:
            try:
                doc = Document()
                
                # إضافة العنوان
                heading = doc.add_heading('النصوص المحددة باللون الأصفر من PDF', 0)
                heading.alignment = 1  # توسيط
                
                # إضافة معلومات الملف
                info_para = doc.add_paragraph()
                info_para.add_run(f"الملف المصدر: {os.path.basename(self.pdf_file)}\n").bold = True
                info_para.add_run(f"تاريخ الاستخراج: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
                info_para.add_run(f"عدد النصوص المستخرجة: {len(self.extracted_highlights)}")
                
                doc.add_paragraph()  # سطر فارغ
                
                # إضافة النصوص المستخرجة
                for i, highlight in enumerate(self.extracted_highlights, 1):
                    # رقم وصفحة النص
                    header_para = doc.add_paragraph()
                    header_para.add_run(f"[{i}] صفحة {highlight['page']}:").bold = True
                    
                    # النص المحدد
                    text_para = doc.add_paragraph(highlight['text'])
                    text_para.style = 'Quote'
                    
                    # خط فاصل
                    doc.add_paragraph('_' * 50)
                
                doc.save(file_path)
                messagebox.showinfo("نجح", f"تم حفظ الملف بنجاح في:\n{file_path}")
                self.status_var.set("تم حفظ ملف Word بنجاح")
                
            except Exception as e:
                messagebox.showerror("خطأ", f"حدث خطأ أثناء حفظ الملف: {str(e)}")
    
    def run(self):
        """تشغيل التطبيق"""
        self.root.mainloop()


def main():
    """الدالة الرئيسية"""
    app = PDFHighlightExtractor()
    app.run()


if __name__ == "__main__":
    main()
